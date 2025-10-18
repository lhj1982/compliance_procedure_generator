from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from werkzeug.utils import secure_filename
from io import BytesIO
from docx import Document
import os
import psycopg2
import psycopg2.extras
from openai import OpenAI
from dotenv import load_dotenv
import logging
import sys
from storage_handler import StorageHandler

load_dotenv()

# Configure logging to stdout (Docker logs)
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    stream=sys.stdout
)
logger = logging.getLogger(__name__)

# Load API key and base URL from environment variables
API_KEY = os.getenv("LLM_API_KEY")
BASE_URL = os.getenv("LLM_BASE_URL")

# Database configuration
DB_CONFIG = {
    'host': os.getenv('DB_HOST', 'localhost'),
    'database': os.getenv('DB_NAME', 'compliance_admin'),
    'user': os.getenv('DB_USER', 'postgres'),
    'password': os.getenv('DB_PASSWORD', 'password'),
    'port': os.getenv('DB_PORT', '5432')
}

client = OpenAI(
    api_key=API_KEY,
    base_url=BASE_URL
)

app = Flask(__name__)
CORS(app)  # Enable CORS for all routes

INITIAL_PROMPT = "You are given brief, informal answers from a subject-matter expert (SME). Your task is to convert those answers into a **formal, auditor-quality standard operating procedure (SOP)**.\n\n" \
"The output must be clear, structured, and repeatable, suitable for internal control, governance, or audit review.\n\n" \
"**Document Template / Structure**\n" \
"Use the following headings (and sub-structure) in every procedure:\n\n" \
"1. Procedure Name\n" \
"2. Owner / Performer (role, team, or individual)\n" \
"3. Frequency (e.g. daily, weekly, monthly, quarterly, ad hoc)\n" \
"4. Purpose / Risk Mitigation\n" \
"5. Procedure Steps (numbered)\n" \
"6. Tools & Systems Used\n" \
"7. Access / Permissions Required\n" \
"8. Starting Point (where work begins)\n" \
"9. Checks & Criteria (standards, thresholds, rules)\n" \
"10. Exception / Failure Handling (escalation, remediation)\n" \
"11. Dependencies / Inputs\n" \
"12. Approvals / Sign-off\n" \
"13. Evidence / Records Storage\n" \
"14. Work Location / Team (onsite, remote, regional)\n" \
"15. Versioning & Review Information (effective date, next review)\n\n" \
"**Language & Style Guidance**\n" \
"- Use formal, compliance-style language: e.g. \"This procedure ensures ...\", \"In the event of failure ...\", \"Escalation is performed to ...\".\n" \
"- If the SME answer is shorthand or partial, expand into clear, full sentences.\n" \
"- Do *not* invent critical facts; if something isn't provided, mark a placeholder (e.g. \"[TBD: Approver]\") rather than guessing.\n" \
"- Maintain numbering consistency and clear hierarchy.\n" \
"- Emphasize **traceability**: each step should map to the checks & criteria, and evidence storage should link to steps.\n\n" \
"**Process**\n" \
"1. You will be given a set of answer pairs: a \"Question\" and \"SME's short answer.\"\n" \
"2. Reformulate into the full procedure document following the template above.\n" \
"3. If any essential information is missing (e.g. approval role), flag it as needing input."

'''
INITIAL_PROMPT = "You are tasked with generating a documented PCI DSS / PCI PIN / P2PE control procedure based on engineers’ answers. " \
"Expand their raw input into a formal, auditor-ready procedure that is clear, traceable, and aligned with compliance standards." \
"Document Structure" \
"The output MUST include these sections:" \
"1.	Control Name – From engineer input." \
"2.	Control Owner / Performer – Roles accountable." \
"3.	Frequency – Daily, weekly, monthly, etc." \
"4.	Scope – Systems, teams, and environments in scope (CDE, AWS, SaaS, etc.)." \
"5.	Purpose / PCI Risk Mitigation – Explicit PCI DSS requirement references (e.g., v4.0 2.4, 7.x, 10.x, 12.x) and how the control reduces risk." \
"6.	Key Definitions – Define CDE, Asset Registry, EOL, LADR, etc." \
"7.	Roles & Responsibilities – RACI-style list of each role and its duties." \
"8.	Procedure Steps – Step-by-step instructions a new engineer can follow. Break down into sub-sections if lifecycle-based (e.g., Intake & Approval, Provisioning, Deployment, Monitoring, Inventory Review, Decommissioning)." \
"9.	Tools & Systems – Platforms used (Jira, GitHub, AWS, Splunk, etc.)." \
"10.	Access Requirements – VPN, SSO, repos, privileged accounts needed." \
"11.	Starting Point – Exact place work begins (dashboard path, Jira board, console link)." \
"12.	Checks & Criteria – What standards or thresholds must be verified (e.g., duplication checks, least privilege, monitoring enabled, EOL patching)." \
"13.	Failure Handling / Escalation – Define clear escalation paths, SLA expectations, and who makes final decisions." \
"14.	Evidence & Recordkeeping – Evidence sources, retention requirements (≥12 months, last 3 months readily available), access controls." \
"15.	Approval / Sign-off – Who approves (hiring manager, IT manager, security)." \
"16.	Enforcement & Exceptions – Non-compliance consequences; exception workflow with risk acceptance and compensating controls." \
"17.	Dependencies – Other controls, teams, or systems this depends on." \
"18.	Document Control – Version, effective date, last review, next review due, change history." \
"Style & Compliance Rules" \
"•	Use formal compliance language:" \
"Example: “The control ensures…”, “Evidence must be retained…”, “Escalation occurs if…”. " \
"•	Always include PCI DSS mapping (e.g., PCI DSS v4.0 2.4, 7.2.1, 10.2.2, 12.3.1)." \
"•	Ensure procedures are measurable and auditable (no vague “should” or “may”; use MUST / SHALL)." \
"•	Include metrics where possible (e.g., 100% of assets inventoried, patch SLA = 30 days)." \
"•	Structure output so it is audit-ready and repeatable. " \
"Put your answer of the section, numbered steps, and bullet points to match the section in uploaded template file, and use the same format to generate the document."
'''

def get_template_from_docx(docx_path):
    doc = Document(docx_path)
    content = []
    for para in doc.paragraphs:
        content.append(para.text)
    template_text = "\n".join(content)
    return template_text

def extract_template_sections(docx_path):
    doc = Document(docx_path)
    sections = []
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            sections.append(para.text)
    return sections

# Get the directory where this script is located
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_TEMPLATE_PATH = os.path.join(SCRIPT_DIR, "Procedure.docx")

@app.route("/", methods=["GET"])
def health_check():
    logger.info("Health check endpoint accessed")
    return jsonify({"status": "healthy", "service": "compliance-procedure-generator-api"})

def create_docx_from_gpt(template_path, gpt_answer):
    template_doc = Document(template_path)
    new_doc = Document()
    answer_sections = {}
    current_section = None
    for line in gpt_answer.split('\n'):
        # Match section headings from template
        if line.strip() in [p.text for p in template_doc.paragraphs if p.style.name.startswith('Heading')]:
            current_section = line.strip()
            answer_sections[current_section] = []
        elif current_section:
            answer_sections[current_section].append(line)
    # Build new docx using template headings and answer content
    for para in template_doc.paragraphs:
        if para.style.name.startswith('Heading'):
            # Use built-in heading style
            level = int(para.style.name.replace('Heading ', ''))
            new_doc.add_heading(para.text, level=level)
            content = answer_sections.get(para.text, [])
            for c in content:
                if c.strip().startswith('- '):
                    new_doc.add_paragraph(c.strip()[2:], style='List Bullet')
                elif c.strip().startswith('1.') or c.strip().startswith('2.'):
                    new_doc.add_paragraph(c.strip(), style='List Number')
                else:
                    new_doc.add_paragraph(c.strip(), style='Normal')
        else:
            # Use 'Normal' for non-heading paragraphs
            new_doc.add_paragraph(para.text, style='Normal')
    return new_doc

@app.route("/download", methods=["POST"])
def download():
    answer = request.form["answer"]
    doc = create_docx_from_gpt("Procedure.docx", answer)
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return send_file(
        file_stream,
        as_attachment=True,
        download_name="procedure_document.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

"""
# Example API endpoint
@app.route('/api/generate_procedure', methods=['POST'])
def generate_procedure():
    answer = ""
    template_prompt = INITIAL_PROMPT + "\n\n" + get_template_from_docx(DOCX_TEMPLATE_PATH)
    user_input = ""
    if request.method == "POST":
        uploaded_file = request.files.get("userfile")
        if uploaded_file:
            filename = secure_filename(uploaded_file.filename)
            if filename.endswith(".txt"):
                user_input = uploaded_file.read().decode("utf-8")
            elif filename.endswith(".docx"):
                doc = Document(uploaded_file)
                user_input = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            else:
                user_input = ""
        if user_input:
            response = client.chat.completions.create(
                model="gpt-5",
                messages=[
                    {"role": "system", "content": template_prompt},
                    {"role": "user", "content": user_input}
                ]
            )
            answer = response.choices[0].message.content
    return {"answer": answer}
"""

def get_db_connection():
    """Get database connection"""
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        return conn
    except psycopg2.Error as e:
        logger.error(f"Database connection error: {e}")
        return None

@app.route('/api/teams', methods=['GET'])
def get_teams():
    """Get all teams from database"""
    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500

    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT id, name FROM teams ORDER BY name")
        teams = cur.fetchall()
        cur.close()
        conn.close()
        logger.info(f"Retrieved {len(teams)} teams from database")
        return jsonify([dict(team) for team in teams])
    except psycopg2.Error as e:
        logger.error(f"Database query error: {e}")
        if conn:
            conn.close()
        return jsonify({'error': 'Failed to fetch teams'}), 500

@app.route('/api/teams/<int:team_id>/questions', methods=['GET'])
def get_team_questions(team_id):
    """Get questions for a specific team"""
    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500

    try:
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT id, name, questions FROM teams WHERE id = %s", (team_id,))
        team = cur.fetchone()
        cur.close()
        conn.close()

        if not team:
            return jsonify({'error': 'Team not found'}), 404

        logger.info(f"Retrieved questions for team: {team['name']}")
        return jsonify({
            'team_id': team['id'],
            'team_name': team['name'],
            'questions': team['questions'] if team['questions'] else []
        })
    except psycopg2.Error as e:
        logger.error(f"Database query error: {e}")
        if conn:
            conn.close()
        return jsonify({'error': 'Failed to fetch team questions'}), 500

@app.route('/api/submit_answers', methods=['POST'])
def submit_answers():
    """Handle form submission and generate compliance document"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400

        team_id = data.get('team_id')
        team_name = data.get('team_name')
        answers = data.get('answers', {})

        if not team_id or not answers:
            return jsonify({'error': 'Team ID and answers are required'}), 400

        # Convert answers to text format for AI processing
        # user_input = f"Team: {team_name}\n\n"
        user_input = ""
        for answer_data in answers.values():
            user_input += f"Q: {answer_data['question']}\n"
            user_input += f"A: {answer_data['answer']}\n\n"

        # logger.info(f"User input for AI:\n{user_input}")
        # Generate document using AI
        template_prompt = INITIAL_PROMPT + "\n\n" + get_template_from_docx(DOCX_TEMPLATE_PATH)

        response = client.chat.completions.create(
            model="gpt-5",
            messages=[
                {"role": "system", "content": template_prompt},
                {"role": "user", "content": user_input}
            ]
        )

        ai_answer = response.choices[0].message.content

        # Create document
        doc = create_docx_from_gpt(DOCX_TEMPLATE_PATH, ai_answer)

        # Save document with team_id naming convention
        document_name = f"{team_id}_procedure_document.docx"

        # Use StorageHandler to save (works with both GCS and local)
        document_path = StorageHandler.save_document(doc, document_name)

        # Save submission to database using upsert logic (insert or update if team already exists)
        conn = get_db_connection()
        if conn:
            try:
                cur = conn.cursor()
                cur.execute("""
                    INSERT INTO teams_compliance_procedures
                    (team_id, document_name, submission_data, status, created_at, updated_at)
                    VALUES (%s, %s, %s, %s, NOW(), NOW())
                    ON CONFLICT (team_id) DO UPDATE SET
                        document_name = EXCLUDED.document_name,
                        submission_data = EXCLUDED.submission_data,
                        status = EXCLUDED.status,
                        updated_at = NOW()
                """, (team_id, document_name, psycopg2.extras.Json(data), 'generated'))
                conn.commit()
                cur.close()
                conn.close()
            except psycopg2.Error as e:
                logger.error(f"Database upsert error: {e}")
                if conn:
                    conn.close()

        # Return success response with download info
        logger.info(f"Successfully generated document: {document_name} for team_id: {team_id}")
        return jsonify({
            'success': True,
            'document_name': document_name,
            'download_url': f'/api/download/{document_name}',
            'message': 'Document generated successfully'
        })

    except Exception as e:
        logger.error(f"Error processing submission: {e}")
        return jsonify({'error': 'Failed to generate document'}), 500

@app.route('/api/download/<filename>', methods=['GET'])
def download_generated_file(filename):
    """Download generated document"""
    try:
        # Security check - ensure filename is safe
        safe_filename = secure_filename(filename)

        # Check if document exists
        if not StorageHandler.document_exists(safe_filename):
            return jsonify({'error': 'File not found'}), 404

        # Retrieve document from storage
        file_stream = StorageHandler.get_document(safe_filename)

        return send_file(
            file_stream,
            as_attachment=True,
            download_name=safe_filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        logger.error(f"Error downloading file: {e}")
        return jsonify({'error': 'Failed to download file'}), 500

if __name__ == "__main__":
    logger.info("Starting Flask application on port 9090")
    app.run(debug=True, host='0.0.0.0', port=9090)